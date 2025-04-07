from docx.shared import Inches
from io import BytesIO
from docx.oxml.ns import nsmap

def overlay_docx(doc, text_map, translated_texts, image_map):
   # Create a lookup for images by OCR text for faster access
   image_lookup = {img['ocr_text']: img for img in image_map}
    
   # Go through each text entry with its translation
   for entry, translated in zip(text_map, translated_texts):
      try:
         if entry["type"] == "paragraph":
            process_paragraph(doc, entry, translated)
                
         elif entry["type"] == "table_cell":
            process_table_cell(doc, entry, translated)
                
         elif entry["type"] == "ocr_image":
            process_image(doc, entry, translated, image_lookup)
                
      except Exception as e:
         print(f"Error processing entry {entry}: {str(e)}")
         continue
         
   return doc

def process_paragraph(doc, entry, translated):
   para = doc.paragraphs[entry["index"]]
   if para.runs:
      first_run = para.runs[0]
      para.clear()
      new_run = para.add_run(translated)
      
      # Copy paragraph formatting (only works if used on entire paragraph)
      new_run.font.name = first_run.font.name
      new_run.font.size = first_run.font.size
      new_run.bold = first_run.bold
      new_run.font.italic = first_run.font.italic
      new_run.font.underline = first_run.font.underline
      new_run.font.color.rgb = first_run.font.color.rgb
   else:
      para.add_run(translated)

def process_table_cell(doc, entry, translated):
   table = doc.tables[entry["table_idx"]]
   cell = table.rows[entry["row"]].cells[entry["col"]]
   para = cell.paragraphs[entry["para_idx"]]
   
   if para.runs:
      first_run = para.runs[0]
      para.clear()
      new_run = para.add_run(translated)
        
      # Copy cell formatting
      new_run.font.name = first_run.font.name
      new_run.font.size = first_run.font.size
      new_run.bold = first_run.bold
      new_run.font.italic = first_run.font.italic
      new_run.font.underline = first_run.font.underline
      new_run.font.color.rgb = first_run.font.color.rgb
   else:
      para.add_run(translated)
   
def process_image(doc, entry, translated, image_lookup):   
   img_entry = image_lookup.get(entry["ocr_text"])
   if not img_entry:
      print(f"Image not found.")
      return

   # Try to find by stored relationship ID (most accurate)
   rel_id = img_entry.get("rel_id")
   if rel_id:
      target_paragraph = find_image_run(doc, rel_id)
      if target_paragraph:
            # Insert caption after the found paragraph
            new_para = doc.add_paragraph(translated, style='Caption')
            target_paragraph._p.addnext(new_para._p)
            print("Added by rID.")
            return
   
   # Add to back if rId fails
   print("Added at end.")
   image_stream = BytesIO(img_entry["image_blob"])
   image_para = doc.add_paragraph()
   image_run = image_para.add_run()
   image_run.add_picture(image_stream, width=Inches(5))
   doc.add_paragraph(translated, style='Caption')
   
def find_image_run(doc, rel_id):
   for paragraph in doc.paragraphs:
      for run in paragraph.runs:
         # Search for both possible element structures
         blip_elements = run._element.xpath('.//*[contains(name(), "blip")]')
         for blip in blip_elements:
            embed_attr = blip.get('{%s}embed' % nsmap['r']) or blip.get('embed')
            if embed_attr == rel_id:
               return paragraph
   return None
